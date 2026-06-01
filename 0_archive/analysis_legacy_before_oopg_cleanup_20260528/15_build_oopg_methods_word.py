"""Build an OOPG-only methods story Word document.

Output:
    6_output/main_output/OOPG_CHE_methods_story.docx

The document is meant for reviewing the paper's method narrative before the
next analysis round. It keeps the current OOPG budget-share results as-is and
explains the planned adult-equivalence capacity-to-pay addition.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = PROJECT_ROOT / "6_output" / "main_output"
CH18_OOPG = OUT_DIR / "chapter18" / "oopg"
DOCX_PATH = OUT_DIR / "OOPG_CHE_methods_story.docx"


def pct(x: float) -> str:
    return f"{100 * float(x):.2f}%"


def num(x: float) -> str:
    return f"{float(x):.3f}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def style_table(table, widths: list[float], header_fill: str = "F2F4F7") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_cell_margins(table)
    set_table_width(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    style_table(table, widths)
    return table


def add_note(doc: Document, title: str, body: str, fill: str = "F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.size = Pt(10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.25])


def add_formula(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_cell_margins(table, top=120, bottom=120, start=180, end=180)
    set_cell_shading(table.cell(0, 0), "FAFAFA")
    p = table.cell(0, 0).paragraphs[0]
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
    set_table_width(table, [6.25])


def add_source(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = "Source"


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 20, "0B2545", 0, 8),
        ("Subtitle", 11, "555555", 0, 12),
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        if name == "Title":
            st.font.bold = True

    if "Source" not in styles:
        source = styles.add_style("Source", 1)
    else:
        source = styles["Source"]
    source.font.name = "Calibri"
    source.font.size = Pt(8)
    source.font.color.rgb = RGBColor(89, 89, 89)
    source.paragraph_format.space_before = Pt(2)
    source.paragraph_format.space_after = Pt(6)


def load_outputs():
    che = pd.read_csv(OUT_DIR / "che_summary.csv")
    che = che[che["scope"] == "oopg"].copy()
    inc = pd.read_csv(CH18_OOPG / "oopg_box18_1_incidence_intensity.csv")
    eq = pd.read_csv(CH18_OOPG / "oopg_box18_2_distribution_sensitive.csv")
    pov = pd.read_csv(OUT_DIR / "poverty_impact.csv")
    pov = pov[pov["scope"] == "oopg"].copy()
    return che, inc, eq, pov


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    che, inc, eq, pov = load_outputs()

    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("OOPG Catastrophic Health Expenditure in Nepal")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "Methods story and results guide for the NLSS IV research paper | "
        f"Generated {datetime.now().strftime('%d %B %Y')}"
    )

    add_note(
        doc,
        "Purpose",
        "This document is a methods-first review note. It explains what each CHE "
        "denominator means, why we keep the simple total-expenditure analysis, "
        "why adult equivalence belongs in the capacity-to-pay module, and how "
        "the current OOPG results should be read while the paper outline is being finalized.",
    )

    doc.add_heading("1. The Paper's Main Object", level=1)
    doc.add_paragraph(
        "The paper now focuses on OOPG: out-of-pocket spending on communicable "
        "disease or injury recorded in NLSS Section 8(B). The household is the "
        "main unit because the survey records health spending at the person-event "
        "level but financial protection is experienced through the household budget. "
        "Individual health payments are therefore summed to the household level."
    )
    add_table(
        doc,
        ["Concept", "Operational choice", "Rationale"],
        [
            [
                "OOPG numerator",
                "Sum of q08_14_i within household",
                "Captures recent general health, communicable disease, or injury spending without mixing in NCD spending.",
            ],
            [
                "Main denominator",
                "Reconstructed nominal consumption plus OOPG",
                "NLSS welfare excludes health spending; adding OOPG back prevents the denominator from omitting the spending category being measured.",
            ],
            [
                "Main threshold",
                "10% of total expenditure",
                "Easy to interpret and comparable to common financial-protection monitoring practice.",
            ],
        ],
        [1.35, 2.1, 2.85],
    )

    doc.add_heading("2. Layer 1: Simple Total-Expenditure CHE", level=1)
    doc.add_paragraph(
        "This is the layman-friendly analysis and should remain the first result. "
        "It asks a plain budget question: what share of a household's monthly "
        "resources went to OOPG? The result is intuitive because a 10% threshold "
        "can be read directly as: one rupee in ten went to general health spending."
    )
    add_formula(
        doc,
        [
            "x_oopg_mo = nominal_total_cons_month + oopg",
            "oopg_sh_tot = oopg / x_oopg_mo",
            "che_oopg_tot10 = 1 if oopg_sh_tot > 0.10",
        ],
    )
    add_note(
        doc,
        "Why adult equivalence is not used here",
        "If both numerator and denominator are divided by the same adult-equivalence scale, the scale cancels out: "
        "(oopg / E) / (x / E) = oopg / x. Using OOPG against only per-adult-equivalent consumption would be inconsistent because it would compare a household-level bill with a one-adult-equivalent denominator.",
        fill="FFF8E6",
    )

    inc_total = inc[inc["denominator"] == "total"]
    rows = []
    for _, row in inc_total.iterrows():
        rows.append(
            [
                f"{int(row['threshold'])}%",
                pct(row["headcount"]),
                pct(row["headcount_se"]),
                pct(row["overshoot"]),
                pct(row["mpo"]),
            ]
        )
    add_table(
        doc,
        ["Threshold", "Headcount", "SE", "Overshoot", "MPO"],
        rows,
        [1.1, 1.25, 0.9, 1.25, 1.25],
    )
    add_source(
        doc,
        "Current OOPG total-expenditure results from chapter18/oopg/oopg_box18_1_incidence_intensity.csv; household survey weights."
    )
    doc.add_paragraph(
        "Interpretation: the 10% headcount is the clean headline. The 5% result "
        "shows lower-intensity financial pressure, while the 20% and higher "
        "thresholds isolate households with more severe budget disruption."
    )

    doc.add_heading("3. Layer 2: Nonfood Expenditure as a Practical Ability-to-Pay Proxy", level=1)
    doc.add_paragraph(
        "The Chapter 18 book also uses OOP as a share of nonfood expenditure. "
        "This is still a budget-share method, but the denominator is narrower. "
        "It treats food as closer to subsistence and asks how large OOPG is "
        "relative to resources outside food consumption."
    )
    add_formula(
        doc,
        [
            "ctp_nonfood_oopg_mo = pcep_nonfood * hhsize / 12 + oopg_real",
            "oopg_sh_nf = oopg_real / ctp_nonfood_oopg_mo",
            "che_oopg_nf40 = 1 if oopg_sh_nf > 0.40",
        ],
    )
    inc_nf = inc[inc["denominator"] == "nonfood"]
    rows = []
    for _, row in inc_nf.iterrows():
        rows.append(
            [
                f"{int(row['threshold'])}%",
                pct(row["headcount"]),
                pct(row["headcount_se"]),
                pct(row["overshoot"]),
                pct(row["mpo"]),
            ]
        )
    add_table(
        doc,
        ["Threshold", "Headcount", "SE", "Overshoot", "MPO"],
        rows,
        [1.1, 1.25, 0.9, 1.25, 1.25],
    )
    doc.add_paragraph(
        "Interpretation: nonfood headcounts should usually be read as a stricter "
        "view of financial stress. The denominator is smaller than total "
        "consumption, so the same OOPG amount can look more burdensome."
    )

    doc.add_heading("4. Layer 3: Adult-Equivalence Capacity to Pay", level=1)
    doc.add_paragraph(
        "The WHO/Xu capacity-to-pay approach is the most welfare-sensitive layer. "
        "It does not merely divide the health bill by total consumption. Instead, "
        "it first protects a household-specific subsistence amount and then "
        "compares OOPG with the resources left after basic needs."
    )
    add_table(
        doc,
        ["Step", "Formula", "Reason"],
        [
            [
                "Equivalent size",
                "E = (A + 0.5K)^0.75",
                "Children are counted as half an adult and larger households receive economies of scale.",
            ],
            [
                "Food per adult equivalent",
                "food_equiv = food_nom_mo / E",
                "Converts household food spending into an adult-equivalent subsistence measure.",
            ],
            [
                "Subsistence line",
                "Weighted mean food_equiv among 45th-55th food-share households",
                "Uses the middle food-share group as a reference for basic food needs.",
            ],
            [
                "Household subsistence",
                "subsistence_hh = subsistence_line * E",
                "Converts the adult-equivalent line back into a household-specific need.",
            ],
            [
                "Capacity to pay",
                "x_oopg_mo - subsistence_hh, or x_oopg_mo - food_nom_mo for very poor households",
                "Protects basic subsistence before measuring health-payment burden.",
            ],
        ],
        [1.25, 2.3, 2.75],
    )
    add_note(
        doc,
        "Important distinction",
        "Adult equivalence is not a new way to shrink the OOPG bill. The OOPG numerator remains the household's actual health payment. Adult equivalence changes the estimated subsistence requirement and therefore the capacity-to-pay denominator.",
    )

    doc.add_heading("5. What the Current OOPG Results Already Show", level=1)
    eq_focus = eq[
        ((eq["denominator"] == "total") & (eq["threshold"].isin([5, 10, 25, 40])))
        | ((eq["denominator"] == "nonfood") & (eq["threshold"].isin([25, 40])))
    ]
    rows = []
    for _, row in eq_focus.iterrows():
        label = "Total" if row["denominator"] == "total" else "Nonfood"
        rows.append(
            [
                label,
                f"{int(row['threshold'])}%",
                pct(row["headcount"]),
                num(row["concentration_headcount"]),
                pct(row["rank_weighted_headcount"]),
            ]
        )
    add_table(
        doc,
        ["Denominator", "Threshold", "Headcount", "CI of H", "Rank-weighted H"],
        rows,
        [1.2, 1.0, 1.15, 1.1, 1.35],
    )
    doc.add_paragraph(
        "Interpretation: concentration indices tell us whether catastrophic "
        "OOPG is more concentrated among poorer or better-off households after "
        "ranking by pre-OOP welfare. Rank-weighted headcounts combine incidence "
        "with distributional concern, giving more weight to burdens among poorer households."
    )

    pov_row = pov[pov["metric"] == "poverty_headcount"].iloc[0]
    add_table(
        doc,
        ["Poverty measure", "Pre-OOPG", "Post-payment", "Change", "People pushed"],
        [
            [
                "Official poverty line",
                pct(pov_row["pre_estimate"]),
                pct(pov_row["post_estimate"]),
                pct(pov_row["difference"]),
                f"{float(pov_row['people_pushed']):,.0f}",
            ]
        ],
        [1.65, 1.2, 1.2, 1.0, 1.3],
    )
    doc.add_paragraph(
        "The poverty analysis keeps the Chapter 19 framing: pcep is post-payment "
        "welfare because NLSS excludes health by construction, so pre-OOPG welfare "
        "is pcep plus annual real per-capita OOPG. We never compute pcep minus OOPG."
    )

    doc.add_heading("6. Recommended Paper Flow", level=1)
    add_table(
        doc,
        ["Paper section", "Story to tell", "Main output"],
        [
            ["Introduction", "Why OOPG can create financial stress in Nepal.", "Research questions"],
            ["Data", "How NLSS IV health payments and consumption are constructed.", "Variable definition table"],
            ["Methods", "Layered CHE framework: total, nonfood, adult-equivalent CTP.", "Formula table"],
            ["Results 1", "How common and intense OOPG CHE is.", "Headcount, overshoot, MPO"],
            ["Results 2", "Where the burden falls in the welfare distribution.", "CI and rank-weighted measures"],
            ["Results 3", "How OOPG changes poverty status.", "Pre/post poverty table"],
            ["Discussion", "What changes when denominator assumptions become more welfare-sensitive.", "Interpretive synthesis"],
        ],
        [1.3, 3.2, 1.75],
    )

    doc.add_heading("7. Sources to Cite", level=1)
    sources = [
        ("O'Donnell et al., Chapter 18", "https://www.worldbank.org/content/dam/Worldbank/document/HDN/Health/HealthEquityCh18.pdf"),
        ("Xu et al. 2003, multicountry CHE analysis", "https://pubmed.ncbi.nlm.nih.gov/12867110/"),
        ("Xu 2005 WHO distribution of health payments method", "https://www.who.int/publications/i/item/EIP-FER-DP.05.2"),
        ("WHO SDG 3.8.2 indicator page", "https://www.who.int/data/gho/data/indicators/indicator-details/GHO/total-population-with-household-expenditures-on-health-greater-than-10-of-total-household-expenditure-or-income-%28sdg-3-8-2%29-%28-%29"),
        ("Koch equivalence-scale paper", "https://academic.oup.com/heapol/article/33/8/966/5070410"),
    ]
    for label, url in sources:
        p = doc.add_paragraph()
        p.style = "Source"
        p.add_run(label + ": ")
        add_hyperlink(p, url, url)

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()

