"""Build an academic-style OOPG CHE methods excerpt.

Output:
    6_output/main_output/OOPG_CHE_academic_excerpt.docx

This document is a manuscript-style excerpt, not a project report. It focuses on
the OOPG analysis, explains the methodological choices in prose, and includes
current OOPG tables where they would naturally appear in a paper draft.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = PROJECT_ROOT / "6_output" / "main_output"
CH18_OOPG = OUT_DIR / "chapter18" / "oopg"
DOCX_PATH = OUT_DIR / "OOPG_CHE_academic_excerpt.docx"


def pct(x: float) -> str:
    return f"{100 * float(x):.2f}"


def pct_se(x: float) -> str:
    return f"{100 * float(x):.2f}"


def f4(x: float) -> str:
    return f"{float(x):.4f}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def style_table(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_cell_margins(table)
    set_table_width(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(9.2)
                    if row_idx == 0:
                        r.bold = True


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)


def add_source(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = "Source"


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table, widths)
    return table


def add_equation(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Cambria Math"
        r.font.size = Pt(10)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(18)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(10.5)
    subtitle.font.color.rgb = RGBColor(85, 85, 85)
    subtitle.paragraph_format.space_after = Pt(14)

    for name, size, color, before, after in [
        ("Heading 1", 14, "2E74B5", 14, 6),
        ("Heading 2", 12, "1F4D78", 10, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    if "Source" not in styles:
        source = styles.add_style("Source", 1)
    else:
        source = styles["Source"]
    source.font.name = "Calibri"
    source.font.size = Pt(8)
    source.font.color.rgb = RGBColor(89, 89, 89)
    source.paragraph_format.space_after = Pt(6)


def load_data():
    inc = pd.read_csv(CH18_OOPG / "oopg_box18_1_incidence_intensity.csv")
    eq = pd.read_csv(CH18_OOPG / "oopg_box18_2_distribution_sensitive.csv")
    pov = pd.read_csv(OUT_DIR / "poverty_impact.csv")
    pov = pov[(pov["scope"] == "oopg") & (pov["metric"] == "poverty_headcount")]
    return inc, eq, pov


def main() -> None:
    inc, eq, pov = load_data()

    doc = Document()
    setup_styles(doc)

    doc.add_paragraph("Methods excerpt: OOPG catastrophic health expenditure in Nepal", style="Title")
    doc.add_paragraph(
        "Draft manuscript text for review. Current estimates use NLSS IV household data; "
        f"document generated {datetime.now().strftime('%d %B %Y')}.",
        style="Subtitle",
    )

    doc.add_heading("Analytical approach", level=1)
    doc.add_paragraph(
        "This study measures catastrophic health expenditure associated with general health "
        "out-of-pocket spending (OOPG) in Nepal using the fourth Nepal Living Standards "
        "Survey. OOPG is defined as household spending reported in Section 8(B) for "
        "communicable disease or injury during the previous 30 days. Individual health "
        "payments are aggregated to the household because catastrophic expenditure is a "
        "household budget event: the relevant question is not only who became ill, but "
        "whether the resulting payment absorbed a substantial share of household resources."
    )
    doc.add_paragraph(
        "The empirical strategy is deliberately layered. First, we present the transparent "
        "budget-share measure in which OOPG is divided by total household expenditure. "
        "This approach follows the catastrophic-payment framework in which the two core "
        "objects are household out-of-pocket payments and a measure of household resources "
        "(O'Donnell et al., 2008). It also aligns with the SDG 3.8.2 convention, under "
        "which catastrophic health spending is monitored as household health expenditure "
        "exceeding 10% or 25% of total household expenditure or income (WHO, 2026). "
        "Second, we retain a nonfood-expenditure denominator, following the Chapter 18 "
        "argument that total expenditure can understate the burden among poorer households "
        "whose resources are largely absorbed by subsistence needs (O'Donnell et al., 2008). "
        "Third, we add a WHO/Xu-style capacity-to-pay specification, in which subsistence "
        "expenditure is estimated before health spending is evaluated against residual "
        "ability to pay (Xu et al., 2003; Xu, 2005)."
    )

    doc.add_heading("Measurement of total-expenditure CHE", level=1)
    doc.add_paragraph(
        "For the main budget-share measure, the numerator is monthly household OOPG. The "
        "denominator is reconstructed nominal monthly consumption with OOPG added back. "
        "This reconstruction is necessary because the NLSS welfare aggregate excludes "
        "health spending by construction; using it without an add-back would compare "
        "health payments with a resource measure that omits the payment category being "
        "studied. Let T_i denote household OOPG and x_i denote reconstructed monthly "
        "nominal consumption inclusive of OOPG. The total-expenditure share is therefore:"
    )
    add_equation(
        doc,
        [
            "x_i = nominal_total_cons_month_i + OOPG_i",
            "s_i^T = OOPG_i / x_i",
            "CHE_i(z) = 1[s_i^T > z].",
        ],
    )
    doc.add_paragraph(
        "Results are shown across a range of thresholds rather than at a single cut-off. "
        "This follows the recommendation in the health-equity measurement literature that "
        "the threshold z is partly normative and should therefore be made visible to the "
        "reader (O'Donnell et al., 2008). The 10% threshold is emphasized because it is "
        "widely used and is easily interpretable as one-tenth of the household budget."
    )

    total = inc[inc["denominator"] == "total"].copy()
    add_caption(doc, "Table 1. Incidence and intensity of OOPG catastrophic payments using total expenditure")
    add_table(
        doc,
        ["Threshold", "Headcount", "Standard error", "Overshoot", "Mean positive overshoot"],
        [
            [
                f"{int(r.threshold)}%",
                pct(r.headcount),
                pct_se(r.headcount_se),
                pct(r.overshoot),
                pct(r.mpo),
            ]
            for r in total.itertuples()
        ],
        [1.0, 1.15, 1.25, 1.15, 1.6],
    )
    add_source(
        doc,
        "Source: Authors' calculations from NLSS IV. Estimates use household survey weights."
    )
    doc.add_paragraph(
        "At the 10% threshold, 7.40% of households incurred catastrophic OOPG spending. "
        "The corresponding overshoot was 0.84% of total expenditure, implying a mean "
        "positive overshoot of 11.35% among households above the threshold. Thus, the "
        "10% result is not simply a count of households just above the line; among those "
        "classified as catastrophic, OOPG spending exceeded the threshold by a meaningful "
        "additional margin."
    )

    doc.add_heading("Nonfood expenditure as an ability-to-pay proxy", level=1)
    doc.add_paragraph(
        "The second denominator narrows the resource base from total expenditure to "
        "nonfood expenditure. Conceptually, this treats food expenditure as closer to "
        "subsistence and asks whether OOPG is large relative to the household's nonfood "
        "budget. This is not identical to the WHO/Xu capacity-to-pay method, but it is "
        "a practical approximation used in the Chapter 18 framework. In our application, "
        "the nonfood denominator is expressed in real terms and OOPG is added back to "
        "maintain consistency between numerator and denominator:"
    )
    add_equation(
        doc,
        [
            "c_i^NF = pcep_nonfood_i * hhsize_i / 12 + OOPG_real_i",
            "s_i^NF = OOPG_real_i / c_i^NF.",
        ],
    )

    nf = inc[inc["denominator"] == "nonfood"].copy()
    add_caption(doc, "Table 2. Incidence and intensity of OOPG catastrophic payments using nonfood expenditure")
    add_table(
        doc,
        ["Threshold", "Headcount", "Standard error", "Overshoot", "Mean positive overshoot"],
        [
            [
                f"{int(r.threshold)}%",
                pct(r.headcount),
                pct_se(r.headcount_se),
                pct(r.overshoot),
                pct(r.mpo),
            ]
            for r in nf.itertuples()
        ],
        [1.0, 1.15, 1.25, 1.15, 1.6],
    )
    add_source(
        doc,
        "Source: Authors' calculations from NLSS IV. Estimates use household survey weights."
    )
    doc.add_paragraph(
        "As expected, the nonfood denominator identifies a different form of vulnerability. "
        "At the 40% nonfood threshold, 2.16% of households were classified as catastrophic. "
        "This estimate should be read as a stricter welfare interpretation rather than a "
        "replacement for the total-expenditure result."
    )

    doc.add_heading("Adult equivalence and the capacity-to-pay extension", level=1)
    doc.add_paragraph(
        "Adult equivalence is introduced only in the capacity-to-pay specification. This is "
        "a substantive methodological choice. If both OOPG and total consumption were "
        "divided by the same adult-equivalence scale, the scale would cancel algebraically: "
        "(OOPG/E)/(x/E) = OOPG/x. Conversely, comparing household OOPG with consumption "
        "per adult equivalent would mix units by placing a household-level payment over a "
        "one-adult-equivalent denominator. For this reason, adult equivalence is not used "
        "to redefine the simple total-expenditure share. Its appropriate role is in the "
        "estimation of subsistence requirements, where household size and composition "
        "alter the resources needed to reach a basic standard of living."
    )
    doc.add_paragraph(
        "Following the logic of the WHO/Xu method, the adult-equivalence scale is used to "
        "construct a household-specific subsistence requirement, which is then subtracted "
        "from total expenditure to obtain capacity to pay. Koch (2018) emphasizes that "
        "equivalence scales enter the WHO method through the poverty line, subsistence "
        "expenditure, and capacity to pay, and that the effect of alternative scales is "
        "ultimately empirical. For this analysis, we use the project scale E_i = "
        "(A_i + 0.5K_i)^0.75, where adults are household members aged 15 years or older "
        "and children are household members below age 15."
    )
    add_caption(doc, "Table 3. Planned adult-equivalence capacity-to-pay construction")
    add_table(
        doc,
        ["Quantity", "Definition", "Interpretation"],
        [
            [
                "Equivalent household size",
                "E_i = (A_i + 0.5K_i)^0.75",
                "Adjusts household size for children and economies of scale.",
            ],
            [
                "Food expenditure",
                "food_nom_mo_i = pcep_food_i * paasche_i * hhsize_i / 12",
                "Nominal monthly household food consumption.",
            ],
            [
                "Equivalized food",
                "food_equiv_i = food_nom_mo_i / E_i",
                "Food spending per adult-equivalent unit.",
            ],
            [
                "Subsistence line",
                "Weighted mean food_equiv_i among the 45th-55th food-share households",
                "Reference food requirement per adult-equivalent unit.",
            ],
            [
                "Household subsistence",
                "subsistence_i = subsistence_line * E_i",
                "Basic expenditure requirement for household i.",
            ],
            [
                "Capacity to pay",
                "x_i - subsistence_i if food_i >= subsistence_i; otherwise x_i - food_i",
                "Resources available after protecting basic needs.",
            ],
            [
                "CTP catastrophic share",
                "OOPG_i / capacity_to_pay_i",
                "WHO/Xu-style financial hardship measure.",
            ],
        ],
        [1.45, 2.65, 2.2],
    )
    add_source(
        doc,
        "Note: CTP estimates will be produced in the next analysis round and kept separate from the existing total and nonfood outputs."
    )

    doc.add_heading("Distribution-sensitive interpretation", level=1)
    doc.add_paragraph(
        "Headcounts and overshoots describe incidence and intensity, but they do not "
        "indicate where catastrophic payments fall in the welfare distribution. Following "
        "the Chapter 18 framework, we therefore estimate concentration indices for the "
        "catastrophic headcount and overshoot and derive rank-weighted measures that give "
        "greater normative weight to burdens concentrated among poorer households "
        "(O'Donnell et al., 2008; Wagstaff and van Doorslaer, 2003)."
    )
    focus = eq[
        ((eq["denominator"] == "total") & (eq["threshold"].isin([10, 25, 40])))
        | ((eq["denominator"] == "nonfood") & (eq["threshold"].isin([25, 40])))
    ].copy()
    add_caption(doc, "Table 4. Distribution-sensitive OOPG catastrophic-payment measures")
    add_table(
        doc,
        ["Denominator", "Threshold", "Headcount", "CI of headcount", "Rank-weighted headcount"],
        [
            [
                "Total" if r.denominator == "total" else "Nonfood",
                f"{int(r.threshold)}%",
                pct(r.headcount),
                f4(r.concentration_headcount),
                pct(r.rank_weighted_headcount),
            ]
            for r in focus.itertuples()
        ],
        [1.2, 1.0, 1.15, 1.35, 1.55],
    )
    add_source(
        doc,
        "Source: Authors' calculations from NLSS IV. Households are ranked by pre-OOP welfare."
    )

    doc.add_heading("Impoverishment framing", level=1)
    doc.add_paragraph(
        "The poverty analysis follows the health-payments framing in which the observed "
        "NLSS welfare measure is treated as post-payment welfare because health spending "
        "is excluded from the official welfare aggregate. Pre-payment welfare is therefore "
        "constructed by adding annual real per-capita OOPG back to pcep. We do not subtract "
        "OOPG from pcep, since doing so would double-count the exclusion of health spending "
        "from the welfare aggregate."
    )
    if not pov.empty:
        r = pov.iloc[0]
        add_caption(doc, "Table 5. OOPG-associated poverty impact")
        add_table(
            doc,
            ["Measure", "Pre-OOPG", "Post-payment", "Change", "People pushed"],
            [[
                "Poverty headcount",
                pct(r.pre_estimate),
                pct(r.post_estimate),
                pct(r.difference),
                f"{float(r.people_pushed):,.0f}",
            ]],
            [1.55, 1.15, 1.15, 1.0, 1.25],
        )
        add_source(
            doc,
            "Source: Authors' calculations from NLSS IV using individual weights."
        )

    doc.add_heading("References", level=1)
    refs = [
        "Koch, S. F. (2018). Catastrophic health payments: does the equivalence scale matter? Health Policy and Planning, 33(8), 966-973.",
        "O'Donnell, O., van Doorslaer, E., Wagstaff, A., & Lindelow, M. (2008). Analyzing Health Equity Using Household Survey Data: A Guide to Techniques and Their Implementation. World Bank.",
        "Wagstaff, A., & van Doorslaer, E. (2003). Catastrophe and impoverishment in paying for health care: with applications to Vietnam 1993-1998. Health Economics, 12, 921-934.",
        "World Health Organization. (2005). Distribution of health payments and catastrophic expenditures: methodology. Geneva: WHO.",
        "World Health Organization. (2026). Global Health Observatory indicator metadata: SDG 3.8.2 catastrophic health spending.",
        "Xu, K., Evans, D. B., Kawabata, K., Zeramdini, R., Klavus, J., & Murray, C. J. L. (2003). Household catastrophic health expenditure: a multicountry analysis. The Lancet, 362(9378), 111-117.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()

