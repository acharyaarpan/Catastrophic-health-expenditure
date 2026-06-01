"""
Build a plain-language review document for the NLSS IV CHE project.

Output:
    6_output/main_output/NLSS_CHE_method_review.docx
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = PROJECT_ROOT / "6_output" / "main_output"
DOCX_PATH = OUT_DIR / "NLSS_CHE_method_review.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D9E2EC") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table, header_fill: str = "1F4E79") -> None:
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.8)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            elif row_idx % 2 == 0:
                set_cell_shading(cell, "F7FAFC")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            if widths:
                cells[i].width = Inches(widths[i])
    style_table(table)
    doc.add_paragraph()
    return table


def add_callout(doc: Document, title: str, body: str, fill: str = "EAF4F4") -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "B7CCD5")
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    run.font.size = Pt(10.5)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.size = Pt(9.5)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def fmt_pct(x: float) -> str:
    return f"{x:.1f}%"


def fmt_pp(x: float) -> str:
    return f"{x:+.2f} pp"


def load_outputs() -> dict[str, pd.DataFrame]:
    return {
        "prevalence": pd.read_csv(OUT_DIR / "che_prevalence_all.csv"),
        "poverty": pd.read_csv(OUT_DIR / "poverty_impact.csv"),
        "regression_status": pd.read_csv(OUT_DIR / "regression_model_status.csv"),
        "equity": pd.read_csv(OUT_DIR / "che_equity_indices.csv"),
    }


def prevalence_value(prev: pd.DataFrame, scope: str, denominator: str, threshold: float, convention: str) -> float:
    hit = prev[
        (prev["weight"] == "ind_wt")
        & (prev["scope"] == scope)
        & (prev["denominator"] == denominator)
        & (prev["threshold"].round(4) == round(threshold, 4))
        & (prev["convention"] == convention)
    ]
    return float(hit.iloc[0]["prevalence"]) * 100.0


def poverty_row(poverty: pd.DataFrame, scope: str) -> pd.Series:
    return poverty[(poverty["scope"] == scope) & (poverty["metric"] == "poverty_headcount")].iloc[0]


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Heading 1", 15, RGBColor(31, 78, 121)),
        ("Heading 2", 12, RGBColor(45, 95, 107)),
        ("Heading 3", 10.5, RGBColor(31, 78, 121)),
    ]:
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = color
        styles[name].font.bold = True
        styles[name].paragraph_format.space_before = Pt(10)
        styles[name].paragraph_format.space_after = Pt(5)

    footer = section.footer.paragraphs[0]
    footer.text = "NLSS IV CHE methods review | Generated 21 May 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)
    return doc


def build_doc() -> None:
    tables = load_outputs()
    prev = tables["prevalence"]
    poverty = tables["poverty"]
    equity = tables["equity"]
    reg_status = tables["regression_status"]

    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("NLSS IV Catastrophic Health Expenditure Analysis")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Book concepts, project rationale, and file-by-file guide")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(90, 90, 90)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared for review before drafting the research paper\n").bold = True
    meta.add_run("Project: D:/Projects/CIH-project/consumption\n")
    meta.add_run("Main methods source: Health Equity book, Chapters 18 and 19, with companion Stata do-files")

    doc.add_paragraph()
    add_callout(
        doc,
        "One-sentence project idea",
        "We are measuring whether household health payments are large enough to threaten living standards, "
        "who is most affected, and how much health spending pushes people below the poverty line in Nepal.",
        fill="E8F1FA",
    )

    doc.add_page_break()

    doc.add_heading("1. Why This Document Exists", level=1)
    doc.add_paragraph(
        "This note is meant to be a bridge between the methods book and our actual NLSS IV analysis files. "
        "It explains the concepts in simple language, records why we made each methodological choice, and "
        "shows where each step is implemented in the project."
    )
    add_bullets(
        doc,
        [
            "Use it to review the whole analysis logic before writing the paper.",
            "Use it to check whether the OOPG and OOP definitions match the research question.",
            "Use it to connect every planned table or figure to a file in the project.",
            "Use it to catch any conceptual issue before we turn results into manuscript text.",
        ],
    )

    doc.add_heading("2. The Research Question in Plain Language", level=1)
    doc.add_paragraph(
        "The central question is not simply whether households spend money on health care. The question is "
        "whether those payments are large relative to the household's ability to pay, whether the burden is "
        "unequally distributed across poorer and richer households, and whether health payments are associated "
        "with households falling below the poverty line."
    )
    add_table(
        doc,
        ["Scope", "What it captures", "Why we keep it"],
        [
            [
                "OOPG",
                "Past-30-day out-of-pocket spending on communicable disease or injury from NLSS Section 8(B).",
                "This is the project-specific narrow scope. It lets us study the direct burden of acute illness and injury spending.",
            ],
            [
                "OOP",
                "OOPG plus annual NCD out-of-pocket spending converted to a monthly amount.",
                "This is closer to a broader household health-payment burden and aligns better with the book's general OOP framing.",
            ],
        ],
        widths=[1.0, 3.0, 3.0],
    )
    add_callout(
        doc,
        "Naming note",
        "Earlier drafts used the shorthand GHE for the communicable subset. We renamed it to OOPG because GHE usually means General Government Health Expenditure in national health accounts. Both of our scopes are household out-of-pocket categories.",
        fill="FFF4E5",
    )

    doc.add_heading("3. What Chapter 18 Is Teaching Us", level=1)
    doc.add_paragraph(
        "Chapter 18 is about catastrophic payments for health care. A payment is called catastrophic when it "
        "crosses a chosen threshold relative to a household resource measure. The chapter does not say that one "
        "threshold is the only correct one. Instead, it shows how to construct indicators, compare thresholds, "
        "measure intensity, and examine whether the burden is concentrated among poorer or richer households."
    )
    add_table(
        doc,
        ["Book idea", "Simple meaning", "Our implementation"],
        [
            [
                "Health-payment share",
                "How large health spending is compared with household resources.",
                "`oopg_sh_tot`, `oop_sh_tot`, `oopg_sh_nf`, `oop_sh_nf` in `2_prep/1_catastrophic.do`.",
            ],
            [
                "Catastrophic headcount",
                "The share of households or people whose health-payment share crosses a threshold.",
                "`che_oopg_tot10`, `che_oop_tot10`, plus 20%, 25%, and 40% variants.",
            ],
            [
                "Overshoot",
                "How far above the threshold the household is, counting zero for households below the threshold.",
                "`over_oopg_*` and `over_oop_*` variables.",
            ],
            [
                "Mean positive overshoot",
                "Among only catastrophic households, how severe the excess burden is.",
                "Computed in `3_analysis/2_catastrophic_measures.do` and saved in `che_summary.csv`.",
            ],
            [
                "Distribution-sensitive measures",
                "Whether catastrophic spending falls more on poorer or richer households.",
                "Concentration indices and rank-weighted measures in `che_equity_indices.csv`.",
            ],
        ],
        widths=[1.55, 2.6, 3.1],
    )

    doc.add_heading("4. What Chapter 19 Is Teaching Us", level=1)
    doc.add_paragraph(
        "Chapter 19 is about health payments and poverty. The basic idea is to compare living standards before "
        "and after health payments. If a household was above the poverty line before payments but below it after "
        "payments, that household is counted as pushed into poverty by health payments."
    )
    add_table(
        doc,
        ["Book idea", "Simple meaning", "Our implementation"],
        [
            [
                "Pre-payment welfare",
                "Living standard before paying for health care.",
                "`pre_oopg = pcep + oopg_pc_ann_real` and `pre_oop = pcep + oop_pc_ann_real`.",
            ],
            [
                "Post-payment welfare",
                "Living standard after health payments.",
                "`post = pcep`, because NLSS IV already excludes health from official welfare.",
            ],
            [
                "Poverty headcount",
                "Percent of people below the poverty line.",
                "`pcep < pline` reproduces official poverty: 20.27%.",
            ],
            [
                "Poverty gap",
                "How far poor households are below the poverty line.",
                "Reported in `6_output/poverty_impact.csv`.",
            ],
            [
                "People pushed into poverty",
                "People whose pre-payment welfare is above the line but post-payment welfare is below it.",
                "Computed for both OOPG and OOP in `3_analysis/3_poverty_impact.do`.",
            ],
        ],
        widths=[1.55, 2.55, 3.15],
    )
    add_callout(
        doc,
        "Core rule for this project",
        "Never compute pcep minus OOP. In NLSS IV, pcep is already health-excluding. For impoverishment we reconstruct pre-payment welfare by adding health payments back to pcep.",
        fill="FCEAEA",
    )

    doc.add_heading("5. The Key Adaptation for NLSS IV", level=1)
    doc.add_paragraph(
        "The book examples use a dataset where total expenditure includes health spending, so the book can "
        "subtract health payments to move from gross to net living standards. NLSS IV is different. The official "
        "welfare aggregate excludes health spending by construction. That changes our implementation."
    )
    add_table(
        doc,
        ["Issue", "If ignored", "Our solution"],
        [
            [
                "NLSS welfare excludes health",
                "The CHE denominator would omit the spending category in the numerator, inflating shares.",
                "Primary denominators add total real OOP back: `totexp_real_hh_mo` and `ctp_real_hh_mo`.",
            ],
            [
                "Poverty analysis needs same units",
                "Subtracting nominal OOP from real pcep would mix units and reverse the logic.",
                "Deflate OOP by `paasche`, convert to real annual per-capita amounts, and add back to pcep.",
            ],
            [
                "Equity ranking should be before payment",
                "Ranking by post-payment pcep can mechanically move high-OOP households downward.",
                "Primary concentration indices rank by `pre_oop`; pcep rank is sensitivity.",
            ],
        ],
        widths=[1.6, 2.65, 3.0],
    )

    doc.add_heading("6. Analysis Pipeline and File Map", level=1)
    add_table(
        doc,
        ["Stage", "File", "What it does"],
        [
            ["Master run", "0_master.do", "Sets folder globals, ado paths, and runs all Stata prep and analysis files."],
            ["Data preparation", "2_prep/1_catastrophic.do", "Merges NLSS files, creates OOPG/OOP variables, denominators, CHE flags, overshoots, and validation assertions."],
            ["Descriptives", "3_analysis/0_descriptive.do", "Creates weighted descriptive statistics and CHE prevalence rows."],
            ["Regression", "3_analysis/1_logit_che.do", "Runs survey-weighted logit models for main and supplementary CHE outcomes."],
            ["CHE and equity", "3_analysis/2_catastrophic_measures.do", "Creates headcount, overshoot, sensitivity, subgroup, CI, and rank-weighted outputs."],
            ["Poverty impact", "3_analysis/3_poverty_impact.do", "Creates Chapter 19 pre/post poverty and impoverishment estimates."],
            ["Figures", "3_analysis/2_pens_parade.py; 3_analysis/3_pens_parade_oop.py", "Creates Pen's Parade figures and the figure audit data workbook."],
            ["Audit", "3_analysis/4_audit_workbook.py", "Creates `audit_workbook.xlsx`, reconciling Stata and Python and documenting derived variables."],
        ],
        widths=[1.25, 2.2, 3.8],
    )

    doc.add_heading("7. Current Headline Results for Review", level=1)
    prevalence_rows = []
    for scope_label, scope in [("OOPG", "oopg"), ("OOP", "oop")]:
        prevalence_rows.append(
            [
                scope_label,
                "Total 10%",
                fmt_pct(prevalence_value(prev, scope, "total", 0.10, "primary")),
                fmt_pct(prevalence_value(prev, scope, "total", 0.10, "nlss")),
            ]
        )
        prevalence_rows.append(
            [
                scope_label,
                "Total 20%",
                fmt_pct(prevalence_value(prev, scope, "total", 0.20, "primary")),
                fmt_pct(prevalence_value(prev, scope, "total", 0.20, "nlss")),
            ]
        )
        prevalence_rows.append(
            [
                scope_label,
                "Nonfood/CTP 25%",
                fmt_pct(prevalence_value(prev, scope, "nonfood", 0.25, "primary")),
                fmt_pct(prevalence_value(prev, scope, "nonfood", 0.25, "nlss")),
            ]
        )
        prevalence_rows.append(
            [
                scope_label,
                "Nonfood/CTP 40%",
                fmt_pct(prevalence_value(prev, scope, "nonfood", 0.40, "primary")),
                fmt_pct(prevalence_value(prev, scope, "nonfood", 0.40, "nlss")),
            ]
        )
    add_table(
        doc,
        ["Scope", "Threshold", "Primary OOPG-addback", "NLSS-style sensitivity"],
        prevalence_rows,
        widths=[0.9, 1.65, 1.55, 1.75],
    )

    pov_rows = []
    for scope_label, scope in [("OOPG", "oopg"), ("OOP", "oop")]:
        row = poverty_row(poverty, scope)
        pov_rows.append(
            [
                scope_label,
                fmt_pct(float(row["pre_estimate"]) * 100),
                fmt_pct(float(row["post_estimate"]) * 100),
                fmt_pp(float(row["difference"]) * 100),
                f"{float(row['people_pushed']):,.0f}",
                f"{int(row['households_pushed']):,}",
            ]
        )
    add_table(
        doc,
        ["Scope", "Pre-payment poor", "Post-payment poor", "Change", "People pushed", "HHs crossing"],
        pov_rows,
        widths=[0.8, 1.35, 1.35, 0.95, 1.25, 1.05],
    )

    doc.add_heading("8. Equity, Regressions, and What They Add", level=1)
    doc.add_paragraph(
        "The descriptive and poverty results show how common the burden is. The equity and regression files add "
        "two further layers: whether the burden is concentrated by welfare rank, and which household characteristics "
        "are associated with the probability of catastrophic spending."
    )
    ci_rows = []
    ci_focus = equity[
        (equity["weight"] == "ind_wt")
        & (equity["measure"] == "headcount")
        & (equity["variable"].isin(["che_oopg_tot10", "che_oop_tot10", "che_oopg_nf40", "che_oop_nf40"]))
    ]
    for _, row in ci_focus.iterrows():
        ci_rows.append(
            [
                row["variable"],
                f"{row['mean'] * 100:.1f}%",
                f"{row['ci_pre_oop']:.3f}",
                f"{row['ci_pcep']:.3f}",
                f"{row['ci_delta']:.3f}",
            ]
        )
    add_table(
        doc,
        ["Outcome", "Mean", "CI pre-OOP rank", "CI pcep rank", "Delta"],
        ci_rows,
        widths=[1.65, 0.8, 1.4, 1.15, 0.8],
    )

    doc.add_page_break()
    doc.add_heading("Regression Model Status", level=2)
    reg_rows = []
    for _, row in reg_status.iterrows():
        reg_rows.append([row["model"], row["outcome"], row["role"], str(row["rc"]), row["note"]])
    add_table(
        doc,
        ["Model", "Outcome", "Role", "RC", "Status note"],
        reg_rows,
        widths=[1.2, 1.35, 0.9, 0.55, 3.2],
    )
    add_callout(
        doc,
        "Interpretation caution",
        "Regression coefficients should be presented as associations, not causal effects. The model helps describe which household profiles have higher CHE risk after adjusting for observed covariates.",
        fill="F4F7FB",
    )

    doc.add_heading("9. What Each Output Is For", level=1)
    add_table(
        doc,
        ["Output", "Use in review or paper"],
        [
            ["descriptive_means.csv", "Table 1 style sample and household profile statistics."],
            ["che_summary.csv", "Overall CHE headcount, overshoot, and mean positive overshoot."],
            ["che_sensitivity.csv", "Primary OOPG-addback estimates compared with NLSS-style denominators."],
            ["che_subgroups.csv", "CHE prevalence by poverty, quintile, province, area, caste/ethnicity, education, elderly member, and disability."],
            ["che_equity_indices.csv", "Concentration indices and rank-weighted measures using pre-OOP rank plus pcep sensitivity."],
            ["poverty_impact.csv", "Pre/post poverty, poverty gap, and people pushed into poverty."],
            ["pens_parade.pdf/png", "Visual distribution of welfare and health-payment shares."],
            ["pens_parade_oop.pdf/png", "Visual pre/post-payment poverty framing."],
            ["audit_workbook.xlsx", "Single workbook for checking variable formulas, Stata/Python reconciliation, sanity checks, and prevalence tables."],
        ],
        widths=[2.0, 5.1],
    )

    doc.add_heading("10. Decisions to Review Before Writing", level=1)
    add_bullets(
        doc,
        [
            "Should the paper emphasize OOP as the main health-financing result, with OOPG as a narrower sensitivity or disease-scope analysis?",
            "Do we want one paper with both scopes, or one primary paper on OOP and a secondary communicable/injury-focused analysis?",
            "Which threshold should be highlighted in the abstract: total 10%, nonfood/CTP 40%, or both?",
            "How should we frame annualizing 30-day communicable/injury spending for poverty analysis?",
            "Do we want to include regression models if supplementary nonfood models have convergence problems?",
            "Should subgroup results be narrowed for the paper to avoid too many tables?",
        ],
    )

    doc.add_page_break()
    doc.add_heading("11. Suggested Paper Flow", level=1)
    add_table(
        doc,
        ["Paper section", "Suggested content"],
        [
            ["Introduction", "Why health payments can create financial hardship in Nepal; why NLSS IV is useful."],
            ["Methods", "Data, OOPG/OOP definitions, reconstructed nominal OOPG-addback denominators, CHE thresholds, concentration indices, poverty framing, survey weights."],
            ["Results 1", "Descriptive profile and OOP spending distribution."],
            ["Results 2", "CHE headcount, overshoot, sensitivity, and subgroup patterns."],
            ["Results 3", "Equity results: concentration indices and rank-weighted measures."],
            ["Results 4", "Poverty and impoverishment results, with OOP emphasized."],
            ["Discussion", "Meaning of OOP versus OOPG results, policy relevance, limitations, and future work."],
        ],
        widths=[1.5, 5.7],
    )

    doc.add_heading("12. Minimal Review Checklist", level=1)
    add_bullets(
        doc,
        [
            "Confirm the two expenditure scopes are conceptually acceptable: OOPG and OOP.",
            "Confirm the primary denominator convention: reconstructed nominal total consumption plus OOPG, with NLSS-style as sensitivity.",
            "Confirm poverty analysis uses add-back logic and never pcep minus health payments.",
            "Confirm the primary CI rank should be pre-OOP, with pcep rank as sensitivity.",
            "Open `audit_workbook.xlsx` and inspect the variable dictionary, trace households, and sanity checks.",
            "Review whether the latest headline numbers are plausible and policy-relevant.",
            "Decide which outputs should become manuscript tables and which should remain appendices.",
        ],
    )

    doc.add_heading("Appendix: Core Variable Names", level=1)
    add_table(
        doc,
        ["Concept", "Variable names"],
        [
            ["Payment scopes", "`oopg`, `oop`, `oopg_ann`, `oop_ann`, `oopg_real`, `oop_real`"],
            ["Per-capita poverty add-back", "`oopg_pc_ann_real`, `oop_pc_ann_real`, `pre_oopg`, `pre_oop`"],
            ["Primary total CHE", "`che_oopg_tot10`, `che_oop_tot10`, `che_oopg_tot20`, `che_oop_tot20`"],
            ["Primary nonfood/CTP CHE", "`che_oopg_nf25`, `che_oop_nf25`, `che_oopg_nf40`, `che_oop_nf40`"],
            ["Sensitivity CHE", "Same CHE names with suffix `_nlss`"],
            ["Overshoot", "`over_oopg_tot10`, `over_oop_tot10`, etc."],
            ["Welfare ranking", "`pre_oop` for primary CI ranking; `pcep` for sensitivity ranking"],
        ],
        widths=[1.8, 5.3],
    )

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    build_doc()
